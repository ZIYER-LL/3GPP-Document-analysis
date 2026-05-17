from fastapi import APIRouter, Depends, Query
from sqlalchemy.orm import Session
from sqlalchemy import select
from datetime import datetime, timezone
from fastapi import APIRouter, Depends, Query, HTTPException
from app.services.downloader import download_document_by_url

from app.services.llm_summary import summarize_document
from app.services.text_extractor import extract_text

from app.core.database import get_db
from app.models.document import Document

from datetime import datetime, timezone
from fastapi import HTTPException
from app.services.model_client import call_summary_model
from app.services.text_extractor import extract_text

from app.services.archive_handler import resolve_analysis_files, choose_primary_file
from app.services.text_extractor import extract_text, build_analysis_text
from app.services.downloader import download_document_by_url
from app.services.model_client import call_summary_model

from fastapi.responses import Response
from io import BytesIO
from docx import Document as DocxDocument
from fastapi.responses import StreamingResponse
router = APIRouter(prefix="/documents", tags=["documents"])


@router.get("")
def list_documents(
    doc_role: str | None = Query(default=None),
    source: str | None = Query(default=None),
    release: str | None = Query(default=None),
    spec: str | None = Query(default=None),
    db: Session = Depends(get_db),
    agenda_item: str | None = Query(default=None),
#    summary_text: Mapped[str | None] = mapped_column(Text, nullable=True),
#    summary_status: Mapped[str | None] = mapped_column(String(50), nullable=True, default="not_started"),
#    summary_error: Mapped[str | None] = mapped_column(Text, nullable=True),
#    summary_updated_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True),
):
    stmt = select(Document)

    if doc_role:
        stmt = stmt.where(Document.doc_role == doc_role)
    if source:
        stmt = stmt.where(Document.source == source)
    if release:
        stmt = stmt.where(Document.release == release)
    if spec:
        stmt = stmt.where(Document.spec == spec)
    if agenda_item:
        stmt = stmt.where(Document.agenda_item == agenda_item)

    result = db.execute(stmt).scalars().all()

    return [
        {
            "id": d.id,
            "tdoc_id": d.tdoc_id,
            "title": d.title,
            "source": d.source,
            "doc_role": d.doc_role,
            "agenda_item": d.agenda_item,
            "agenda_item_desc": d.agenda_item_desc,
            "release": d.release,
            "spec": d.spec,
            "is_cr": d.is_cr,
            "tdoc_url": d.tdoc_url,
            "summary_text": d.summary_text,
            "summary_status": d.summary_status,
            "summary_error": d.summary_error,
        }
        for d in result
    ]


@router.get("/{doc_id}")
def get_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        return {"error": "not found"}

    return {
        "id": doc.id,
        "tdoc_id": doc.tdoc_id,
        "tdoc_url": doc.tdoc_url,
        "title": doc.title,
        "source": doc.source,
        "contact": doc.contact,
        "type_raw": doc.type_raw,
        "for_raw": doc.for_raw,
        "abstract": doc.abstract,
        "agenda": doc.agenda,
        "agenda_item": doc.agenda_item,
        "agenda_item_desc": doc.agenda_item_desc,
        "tdoc_status": doc.tdoc_status,
        "release": doc.release,
        "spec": doc.spec,
        "version": doc.version,
        "related_wi": doc.related_wi,
        "cr": doc.cr,
        "cr_rev": doc.cr_rev,
        "reply_to": doc.reply_to,
        "doc_role": doc.doc_role,
        "is_cr": doc.is_cr,
        "confidence": doc.confidence,
    }

@router.post("/{doc_id}/analyze")
def analyze_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="document not found")

    try:
        doc.summary_status = "processing"
        doc.summary_error = None
        db.commit()

        if not doc.tdoc_url:
            raise HTTPException(
                status_code=400,
                detail="该文稿没有原始链接，当前无法下载正文进行分析"
            )

        local_file_path = download_document_by_url(doc.tdoc_url)

        analysis_files, temp_dir = resolve_analysis_files(local_file_path)

        if not analysis_files:
            raise ValueError("压缩包中未找到可分析的 pdf/docx/txt/md 文件")

        # 方案 A：只选主文件
        primary_file = choose_primary_file(analysis_files)
        if primary_file is None:
            raise ValueError("未找到合适的主文件进行分析")

        # 单文件分析
        # text = extract_text(str(primary_file))

        # 方案 B：多文件拼接分析（推荐）
        text = build_analysis_text(analysis_files)

        if not text or len(text.strip()) < 50:
            raise ValueError("提取到的正文过短，无法生成摘要")

        summary = call_summary_model(
            metadata={
                "tdoc_id": doc.tdoc_id,
                "title": doc.title,
                "source": doc.source,
                "agenda_item": doc.agenda_item,
                "spec": doc.spec,
                "release": doc.release,
            },
            text=text,
        )

        doc.summary_text = summary
        doc.summary_status = "done"
        doc.summary_updated_at = datetime.now(timezone.utc)
        db.commit()

        shutil.rmtree(temp_dir, ignore_errors=True)

        return {
            "document_id": doc.id,
            "analysis_status": doc.summary_status,
            "summary_text": doc.summary_text,
        }

    except HTTPException:
        doc.summary_status = "failed"
        db.commit()
        raise

    except Exception as e:
        doc.summary_status = "failed"
        doc.summary_error = str(e)
        doc.summary_updated_at = datetime.now(timezone.utc)
        db.commit()
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{doc_id}/analysis")
def get_document_analysis(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="document not found")

    return {
        "document_id": doc.id,
        "analysis_status": doc.summary_status,
        "summary_text": doc.summary_text,
        "summary_error": doc.summary_error,
        "updated_at": doc.summary_updated_at,
    }

@router.get("/{doc_id}/analysis/download")
def download_analysis(doc_id: int, format: str = "md", db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="document not found")

    if not doc.summary_text:
        raise HTTPException(status_code=400, detail="暂无分析结果")

    if format == "md":
        content = f"""# {doc.tdoc_id or "3GPP Document"}

        ## 标题
        {doc.title or "-"}

        ## 来源
        {doc.source or "-"}

        ## Agenda Item
        {doc.agenda_item or "-"}

        ## Spec
        {doc.spec or "-"}

        ## Release
        {doc.release or "-"}

        ## AI 摘要
        {doc.summary_text}
        """
        filename = f"{doc.tdoc_id or doc.id}_analysis.md"
        return Response(
            content=content,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    if format == "docx":
        file_obj = BytesIO()
        d = DocxDocument()
        d.add_heading(doc.tdoc_id or "3GPP Document", level=1)
        d.add_paragraph(f"标题：{doc.title or '-'}")
        d.add_paragraph(f"来源：{doc.source or '-'}")
        d.add_paragraph(f"Agenda Item：{doc.agenda_item or '-'}")
        d.add_paragraph(f"Spec：{doc.spec or '-'}")
        d.add_paragraph(f"Release：{doc.release or '-'}")
        d.add_heading("AI 摘要", level=2)
        d.add_paragraph(doc.summary_text)
        d.save(file_obj)
        file_obj.seek(0)

        filename = f"{doc.tdoc_id or doc.id}_analysis.docx"
        return StreamingResponse(
            file_obj,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    raise HTTPException(status_code=400, detail="unsupported format")