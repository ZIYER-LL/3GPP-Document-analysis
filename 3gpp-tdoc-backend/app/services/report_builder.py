from pathlib import Path
from typing import Iterable, Optional

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None


BACKEND_ROOT = Path(__file__).resolve().parents[2]
REPORT_ROOT = BACKEND_ROOT / "reports"


def ensure_job_report_dir(job_id: int) -> Path:
    job_dir = REPORT_ROOT / f"job_{job_id}"
    job_dir.mkdir(parents=True, exist_ok=True)
    return job_dir


def write_markdown_report(path: Path, content: str) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return str(path)


def write_docx_report(path: Path, title: str, paragraphs: Iterable[str]) -> Optional[str]:
    if DocxDocument is None:
        return None

    path.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()
    doc.add_heading(title, level=1)

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    doc.save(str(path))
    return str(path)


def build_item_report_markdown(
    title: str,
    tdoc_id: str | None,
    agenda_item: str | None,
    summary_text: str,
) -> str:
    return f"""# 单篇文稿分析报告

## 基本信息
- 标题：{title}
- TDoc ID：{tdoc_id or "-"}
- Agenda Item：{agenda_item or "-"}

## 摘要
{summary_text}
"""


def build_final_report_markdown(
    meeting_list: str,
    agenda_item: str,
    total_items: int,
    completed_items: int,
    failed_items: int,
    item_sections: list[dict],
) -> str:
    lines = [
        "# Agenda 文稿汇总报告",
        "",
        "## 任务信息",
        f"- Meeting List：{meeting_list}",
        f"- Agenda Item：{agenda_item}",
        f"- 总文稿数：{total_items}",
        f"- 成功完成：{completed_items}",
        f"- 失败数量：{failed_items}",
        "",
        "## 逐篇摘要",
        "",
    ]

    if not item_sections:
        lines.append("暂无可汇总的单篇摘要。")
    else:
        for idx, item in enumerate(item_sections, start=1):
            lines.extend(
                [
                    f"### {idx}. {item.get('title', '-')}",
                    f"- TDoc ID：{item.get('tdoc_id') or '-'}",
                    f"- 状态：{item.get('status') or '-'}",
                    "",
                    item.get("summary_text") or "暂无摘要",
                    "",
                ]
            )

    return "\n".join(lines)